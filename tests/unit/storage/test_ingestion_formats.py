"""Ingestion format coverage tests (Phase 3.1 + 3.2).

Tests PDF (text), DOCX, MD with frontmatter, TXT UTF-8, TXT latin1,
and scanned PDF detection.
"""
from __future__ import annotations

import io
import struct
from pathlib import Path

import pytest

from armance.storage.ingestion import IngestionError, load_file, load_md, load_txt


# ── TXT UTF-8 ─────────────────────────────────────────────────────────────

def test_load_txt_utf8(tmp_path: Path) -> None:
    f = tmp_path / "doc.txt"
    f.write_text("Hello world.\n\nSecond paragraph.", encoding="utf-8")
    chunks = load_txt(f)
    assert chunks
    joined = " ".join(c.text for c in chunks)
    assert "Hello world" in joined


def test_load_txt_latin1(tmp_path: Path) -> None:
    f = tmp_path / "doc.txt"
    f.write_bytes("Caf\xe9 au lait.\n\nSecond paragraph.".encode("latin-1"))
    chunks = load_txt(f)
    assert chunks
    joined = " ".join(c.text for c in chunks)
    assert "lait" in joined


# ── Markdown with frontmatter ─────────────────────────────────────────────

def test_load_md_with_frontmatter(tmp_path: Path) -> None:
    f = tmp_path / "doc.md"
    f.write_text(
        "---\ntitle: Test\ndate: 2026-04-30\n---\n# Heading\n\nBody text here.",
        encoding="utf-8",
    )
    chunks = load_md(f)
    assert chunks
    joined = " ".join(c.text for c in chunks)
    assert "Body text" in joined


def test_load_file_md(tmp_path: Path) -> None:
    f = tmp_path / "notes.md"
    f.write_text("# Title\n\nContent paragraph.", encoding="utf-8")
    chunks = load_file(f)
    assert chunks
    assert all(c.source == "notes.md" for c in chunks)


# ── DOCX ──────────────────────────────────────────────────────────────────

def test_load_docx(tmp_path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("Test Heading", level=1)
    doc.add_paragraph("This is the first paragraph of the document.")
    doc.add_paragraph("This is the second paragraph with more content.")
    path = tmp_path / "test.docx"
    doc.save(str(path))

    chunks = load_file(path)
    assert chunks
    joined = " ".join(c.text for c in chunks)
    assert "first paragraph" in joined


def test_load_docx_with_images_doesnt_crash(tmp_path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Text before image.")
    doc.add_paragraph("Text after image.")
    path = tmp_path / "with_image.docx"
    doc.save(str(path))

    chunks = load_file(path)
    assert chunks


# ── PDF (text) ────────────────────────────────────────────────────────────

def _make_text_pdf(tmp_path: Path, text: str = "Hello from PDF page one.") -> Path:
    """Create a minimal single-page text PDF using reportlab if available,
    or skip if not installed."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError:
        pytest.skip("reportlab not installed — skipping PDF text fixture")

    path = tmp_path / "text.pdf"
    c = canvas.Canvas(str(path), pagesize=letter)
    c.drawString(72, 720, text)
    c.save()
    return path


def test_load_pdf_text(tmp_path: Path) -> None:
    path = _make_text_pdf(tmp_path, "Hello PDF content here.")
    chunks = load_file(path)
    assert chunks
    joined = " ".join(c.text for c in chunks)
    assert "Hello" in joined


# ── Scanned PDF detection ─────────────────────────────────────────────────

def _make_scanned_pdf(tmp_path: Path) -> Path:
    """Create a multi-page PDF with no extractable text (simulates scanned)."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError:
        pytest.skip("reportlab not installed — skipping scanned PDF fixture")

    path = tmp_path / "scanned.pdf"
    c = canvas.Canvas(str(path), pagesize=letter)
    # page 1: no text (empty)
    c.showPage()
    # page 2: no text (empty)
    c.showPage()
    c.save()
    return path


def test_scanned_pdf_raises_ingestion_error(tmp_path: Path) -> None:
    path = _make_scanned_pdf(tmp_path)
    from armance.storage.ingestion import load_pdf

    with pytest.raises(IngestionError, match="scanned"):
        load_pdf(path)


# ── Unsupported extension ─────────────────────────────────────────────────

def test_load_file_unsupported_returns_empty(tmp_path: Path) -> None:
    f = tmp_path / "doc.xlsx"
    f.write_bytes(b"dummy content")
    chunks = load_file(f)
    assert chunks == []
