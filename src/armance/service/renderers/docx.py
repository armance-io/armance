"""DOCX renderer — headings preserved, claim blocks → footnotes.

Requires python-docx (optional [render] extra). Falls back gracefully.

Spec: docs/spec/22_circular_outputs.md § Supported formats (docx)
"""
from __future__ import annotations

import logging
import re
from pathlib import Path

from .base import Deliverable, RenderResult

logger = logging.getLogger(__name__)


class DocxRenderer:
    """Renders deliverables into a Word document."""

    format = "docx"

    async def render(
        self,
        deliverables: list[Deliverable],
        template: Path | None,
        options: dict,
        output_path: Path,
    ) -> RenderResult:
        try:
            from docx import Document  # type: ignore
        except ImportError:
            return RenderResult(
                output_path=output_path,
                warnings=["python-docx not installed; install with: pip install armance[render]"],
                error="python-docx not installed",
            )

        doc = Document(str(template)) if template and template.exists() else Document()
        combined = "\n\n".join(d.content for d in deliverables)
        page_count = 0

        for line in combined.splitlines():
            if line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=1)
                page_count += 1
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
            elif re.match(r"\[\[c_\w+\]\]", line):
                doc.add_paragraph(f"[footnote: {line}]", style="Caption")
            elif line.strip():
                doc.add_paragraph(line.strip())

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        size = output_path.stat().st_size
        return RenderResult(
            output_path=output_path,
            bytes_written=size,
            pages_or_slides=max(page_count, 1),
        )
